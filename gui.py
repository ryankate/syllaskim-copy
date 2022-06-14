from tkinter import *
from tkinter import font
from tkinter import filedialog
from tkinter import messagebox
from turtle import left
import DualInput
from docx import Document
import pandas as pd
import ICSOut

def move(dirn):
    global current

    idx = pages.index(current) + dirn

    if not 0 <= idx < len(pages):
        return

    current.pack_forget()
    current = pages[idx]
    current.pack(side=TOP)

    #if idx == 0:
    #    clean(buttonFrame)
    #    nextButton = Button(buttonFrame, text='Next', command=next)
    #    nextButton.pack(side=RIGHT)
    #    buttonFrame.pack(side=BOTTOM)

    if idx == 0:
        clean(buttonFrame)
        #backButton = Button(buttonFrame, text='Previous', command=prev)
        #backButton.pack(side=LEFT)
        skimButton = Button(buttonFrame, text='Skim', command=process, width=10, height=2, font=buttonFont)
        skimButton.pack(side=RIGHT)
        buttonFrame.pack(side=BOTTOM)

    if idx == 1:
        clean(buttonFrame)
        #toGoogle = Button(buttonFrame, text='Upload to Google Calendar', command=upload)
        #toGoogle.pack(side=TOP)
        calandarExport = Button(buttonFrame, text='Export as .ics file', command=export, width=30, height=2, font=buttonFont)
        calandarExport.pack(side=TOP)
        buttonFrame.pack(side=BOTTOM)

    if idx == 2:
        clean(buttonFrame)
        finalize = Button(buttonFrame, text='Finish', command=quit, width=10, height=2, font=buttonFont)
        finalize.pack(side=LEFT)
        buttonFrame.pack(side=BOTTOM)


def clean(frame):
    list_elements = frame.slaves()
    for l in list_elements:
        l.destroy()


def next():
    move(+1)

def prev():
    move(-1)

def process():
    if filename == "":
        messagebox.showerror("Error", "Please select a file")
    else: 
        global datesList
        # extract the assignments table from the file
        document = Document(filename)
        
        # extract data from each table in document
        for i in range(len(document.tables)):
            table = document.tables[i]
            data = []
            keys = None
            rowdata = None
            
            # get the heading and row info then put it in a data frame
            for j, row in enumerate(table.rows):
                text = (cell.text for cell in row.cells)
                if j == 0:
                    keys = tuple(text)
                    continue
                rowdata = dict(zip(keys, text))
                data.append(rowdata)
            df = pd.DataFrame(data)
            
            #convert from objects to strings
            df = df.convert_dtypes()
            
            #get the column names to make accessing the data frame cleaner
            cols = df.columns
            
            #convert second column from strings to datetime objects (in case that's better for the ics doc)
            df[df.columns[1]] = pd.to_datetime(df[df.columns[1]])
            
            #convert from data frame to list of tuples (since that's how Nathaniel had it)
            datesList = list(zip(df[cols[0]],df[cols[1]]))
        
        # load next page
        page_2.load() 
        move(+1)

def upload():
    # TODO Google Calendar API
    move(+1)

def export():
    f = filedialog.asksaveasfile(mode='w', defaultextension=".ics", title="ICS Save As...", filetypes=(("Calendar File", "*.ics*"),), initialfile="out.ics")
    if f is None: # asksaveasfile return `None` if dialog closed with "cancel".
        return
    ICSOut.saveToICS(datesList, f)
    move(+1)

def browse(fileText):
    global filename
    filename = filedialog.askopenfilename(initialdir = "/",
                                          title = "Select a File",
                                          filetypes = (("Word Files",
                                                        "*.docx*"),
                                                       ("All Files",
                                                        "*.*")))
      
    # Change label contents
    fileText.configure(text="File Opened: "+filename)
    
tk = Tk()

titleFont = ("Times New Roman", 28, "bold")
h1Font = ("Times New Roman", 20, "")
h2Font = ("Times New Roman", 16, "")
h3Font = ("Times New Roman", 12, "")
listboxFont = ("Times New Roman", 14, "")
buttonFont = ("Times New Roman", 14, "")

buttonFrame = Frame(tk)

class page1(Frame):
    def __init__(self, master=None):
        Frame.__init__(self, master)
        self.master = master
        self.pack(fill=BOTH, expand=1)
        
        welcome = Label(self, text='Welcome to SyllaSkim!')
        welcome.configure(font=titleFont)
        welcome.pack() 

        self.img = PhotoImage(file='Syllaskim.png')
        imageLabel = Label(self, image=self.img)
        imageLabel.pack()

        secondary = Label(self, text='This simple wizard will help convert \nyour Syllabus to a list of dates!\n\n', wraplength=0)
        secondary.configure(font=h2Font)
        secondary.pack()


        instruction = Label(self, text='Choose the syllabus file', wraplength=0)
        instruction.configure(font=h1Font)
        instruction.pack(side=TOP)

        fileText = Label(self, text='File:', wraplength=0)
        fileText.configure(font=h3Font)
        fileText.pack(side=TOP)

        skimButton = Button(buttonFrame, text='Skim', command=process, width=10, height=2, font=buttonFont)
        skimButton.pack(side=RIGHT)
        buttonFrame.pack(side=BOTTOM)

        browseButton = Button(self, text='Browse', command=lambda: browse(fileText), width=10, height=1, font=buttonFont)
        browseButton.pack(side=TOP)

        self.pack(side=TOP)

#class page2(Frame):
#    def __init__(self, master=None):
#        Frame.__init__(self, master)
#        self.master = master
#        
#        instruction = Label(self, text='Choose the syllabus file', wraplength=0)
#        instruction.configure(font=h1Font)
#        instruction.pack(side=TOP)
#
#        fileText = Label(self, text='File:', wraplength=0)
#        fileText.configure(font=h3Font)
#        fileText.pack(side=TOP)
#
#        browseButton = Button(self, text='Browse', command=lambda: browse(fileText))
#        browseButton.pack(side=TOP)

class page2(Frame):
    dateListbox = ""

    def __init__(self, master=None):
        Frame.__init__(self, master)
        self.master = master
        
        instruction = Label(self, text='Review the skimmed dates!', wraplength=0)
        instruction.configure(font=h1Font)
        instruction.pack(side=TOP)

        self.dateListbox = Listbox(self, height=10, width=10, font=listboxFont)

        scroll = Scrollbar(self, orient=VERTICAL)
        scroll.pack(side=RIGHT, fill=Y)

        self.dateListbox.configure(yscrollcommand=scroll.set)
        scroll.config(command=self.dateListbox.yview)
        self.dateListbox.pack(fill=X)

        buttonFrame = Frame(self)

        deleteButton = Button(buttonFrame, text="Delete", command=lambda: self.delete(), width=10, height=1, font=buttonFont)
        deleteButton.pack(side=LEFT)

        editButton = Button(buttonFrame, text="Edit", command=lambda: self.edit(), width=10, height=1, font=buttonFont)
        editButton.pack(side=RIGHT)

        addButton = Button(buttonFrame, text="Add...", command=lambda: self.add(), width=10, height=1, font=buttonFont)
        addButton.pack(side=RIGHT)

        buttonFrame.pack(side=TOP)

    def load(self):
        global datesList
        for i in range(len(datesList)):
            toInsert = datesList[i][0] + "   -   " + str(datesList[i][1])
            self.dateListbox.insert(i, toInsert)

    def delete(self):
        global datesList
        index = self.dateListbox.curselection()
        print("Index", index)
        if index != ():
            self.dateListbox.delete(index[0])
            datesList.pop(index[0])
            print(datesList)

    def add(self):
        global datesList
        dual = DualInput.DualInput(self, "Input")
        output = (dual.assignment, pd.to_datetime(dual.date))

        try:
            dual.destroy()
        except:
            pass
        
        if output[0] != None:
            toInsert = output[0] + "   -   " + str(output[1])
            self.dateListbox.insert(self.dateListbox.size(), toInsert)
            datesList.append(output)
            print(datesList)

    def edit(self):
        global datesList
        index = self.dateListbox.curselection()

        if index[0] != None:
            index = index[0]
            dual = DualInput.DualInput(self, "Input", assignment=datesList[index][0], date=datesList[index][1])
            output = (dual.assignment, pd.to_datetime(dual.date))

            try:
                dual.destroy()
            except:
                pass

            edited = output[0] + "   -   " + str(output[1])
            if index != ():
                self.dateListbox.delete(index)
                self.dateListbox.insert(index, edited)
                datesList[index] = output
                print(datesList)

        

class page3(Frame):
    def __init__(self, master=None):
        Frame.__init__(self, master)
        self.master = master
        
        finish = Label(self, text='Task Completed Successfully!', wraplength=0)
        finish.configure(font=h1Font)
        finish.pack(side=TOP)



page_1 = page1(tk)
page_2 = page2(tk)
page_3 = page3(tk)
# page_4 = page4(tk)

pages = [page_1, page_2, page_3]
current = page_1

filename = ""

datesList = []

#nextButton = Button(tk, text='Next', command=next)
#backButton = Button(tk, text='Previous', command=prev)

tk.title("SyllaSkim")

tk.geometry("400x600")

tk.resizable(False, False)

if __name__ == '__main__':
    tk.mainloop()